from __future__ import annotations

import csv
import hashlib
import json
import re
from datetime import date, datetime
from io import BytesIO, StringIO
from typing import Any

import openpyxl
from sqlalchemy.orm import Session

from .agent_service import DEEPSEEK_MODEL, deepseek_json_object
from .execution_importer import find_media_identity_matches, set_with_undo
from .media_taxonomy import normalize_media_payload
from .models import AgentRun, Campaign, CampaignStageEvent, Contact, CostItem, Deliverable, ImportBatch, Media, Product, Project, ProjectProduct, Shipment, ShipmentItem, User
from .product_backfill import ensure_project_link, find_or_create_product, find_product_matches


FIELDS: list[tuple[str, str]] = [
    ("media_name", "媒体名称*"), ("country", "国家"), ("platform", "渠道"), ("profile_url", "主页链接"),
    ("category", "媒体分类"), ("followers_k", "粉丝或流量(K)"), ("contact_name", "联系人姓名"),
    ("contact_role", "联系人职位"), ("contact_email", "联系人邮箱"), ("contact_phone", "联系人电话"),
    ("cooperation_status", "合作状态"), ("project_code", "项目编号"), ("project_name", "项目名称"),
    ("owner_email", "负责人邮箱"), ("collaboration_type", "推广形式"), ("execution_status", "执行状态"),
    ("expected_publish_date", "预计发布日期"), ("next_action", "下一步行动"), ("follow_up_date", "跟进日期"),
    ("product_models", "产品型号（多个用；分隔）"), ("oa_pi_number", "OA/PI"), ("carrier", "承运商"),
    ("tracking_number", "物流单号"), ("shipped_at", "发货日期"), ("delivered_at", "签收日期"),
    ("content_url", "内容链接"), ("published_at", "内容发布日期"), ("planned_cost", "计划费用"),
    ("actual_cost", "实际费用"), ("currency", "币种"), ("payment_status", "付款状态"), ("notes", "备注"),
]
FIELD_KEYS = {key for key, _ in FIELDS}
HEADER_TO_KEY = {label: key for key, label in FIELDS}
STANDARD_HEADERS = [label for _, label in FIELDS]
EXECUTION_STATUSES = {"待确认", "待发货", "运输中", "已签收待产出", "内容审核中", "已发布", "已结算", "已暂停", "已取消"}
DATE_KEYS = {"expected_publish_date", "follow_up_date", "shipped_at", "delivered_at", "published_at"}
NUMBER_KEYS = {"followers_k", "planned_cost", "actual_cost"}
DOCUMENT_EXTENSIONS = {"pdf", "docx", "txt", "md"}
TABLE_EXTENSIONS = {"csv", "xlsx"}
MAX_ROWS = 2000


def template_csv_bytes() -> bytes:
    output = StringIO()
    csv.writer(output, lineterminator="\n").writerow(STANDARD_HEADERS)
    return ("\ufeff" + output.getvalue()).encode("utf-8")


def _text(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    result = str(value).strip()
    return result or None


def _date(value: Any) -> str | None:
    raw = _text(value)
    if not raw:
        return None
    raw = raw.replace("/", "-").replace(".", "-")
    for fmt in ("%Y-%m-%d", "%m-%d-%Y", "%d-%m-%Y", "%Y-%m"):
        try:
            return datetime.strptime(raw, fmt).date().isoformat()
        except ValueError:
            continue
    return raw


def _number(value: Any) -> float | None:
    raw = _text(value)
    if not raw:
        return None
    match = re.search(r"-?[\d,.]+", raw.replace("，", ","))
    if not match:
        return None
    try:
        return float(match.group(0).replace(",", ""))
    except ValueError:
        return None


def _load_table(content: bytes, extension: str) -> tuple[list[str], list[dict[str, Any]]]:
    if extension == "xlsx":
        workbook = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
        sheet = workbook[workbook.sheetnames[0]]
        values = sheet.iter_rows(values_only=True)
        headers = [_text(cell) or "" for cell in next(values, [])]
        rows = [dict(zip(headers, row)) | {"_row_number": index} for index, row in enumerate(values, 2) if any(cell not in (None, "") for cell in row)]
    else:
        decoded = None
        for encoding in ("utf-8-sig", "gb18030", "utf-16"):
            try:
                decoded = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if decoded is None:
            raise ValueError("CSV 编码无法识别，请保存为 UTF-8 CSV")
        reader = csv.DictReader(StringIO(decoded))
        headers = [str(item or "").strip() for item in (reader.fieldnames or [])]
        rows = [dict(row) | {"_row_number": index} for index, row in enumerate(reader, 2) if any(_text(value) for value in row.values())]
    if not headers:
        raise ValueError("文件没有可识别的表头")
    if len(rows) > MAX_ROWS:
        raise ValueError(f"单次最多导入 {MAX_ROWS} 行，请拆分文件")
    return headers, rows


def _extract_document(content: bytes, extension: str) -> str:
    if extension in {"txt", "md"}:
        for encoding in ("utf-8-sig", "gb18030", "utf-16"):
            try:
                return content.decode(encoding)[:45_000]
            except UnicodeDecodeError:
                continue
        raise ValueError("文本编码无法识别")
    if extension == "pdf":
        from pypdf import PdfReader
        text = "\n\n".join((page.extract_text() or "") for page in PdfReader(BytesIO(content)).pages)
    else:
        from docx import Document
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    text = text.strip()
    if not text:
        raise ValueError("文档中没有可提取的文字；扫描版 PDF 请先做 OCR")
    return text[:45_000]


def _agent_table_mapping(headers: list[str], rows: list[dict[str, Any]]) -> tuple[dict[str, str], dict[str, Any]]:
    field_help = {key: label for key, label in FIELDS}
    samples = [{header: _text(row.get(header)) for header in headers} for row in rows[:5]]
    system = """你是 CRM 导入列映射器。只输出 JSON，不改写数据，不猜测不存在的列。一列最多映射到一个标准字段。无法判断则不要映射。输出 {\"mapping\":{\"原表头\":\"标准字段key\"},\"warnings\":[]}。"""
    result, usage = deepseek_json_object(system, json.dumps({"standard_fields": field_help, "headers": headers, "sample_rows": samples}, ensure_ascii=False), max_tokens=2200)
    raw = result.get("mapping") if isinstance(result.get("mapping"), dict) else {}
    mapping = {str(header): str(key) for header, key in raw.items() if header in headers and key in FIELD_KEYS}
    return mapping, {"usage": usage, "warnings": [str(item)[:240] for item in (result.get("warnings") or [])[:20]]}


def _agent_document_records(text: str, filename: str) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    field_help = {key: label for key, label in FIELDS}
    system = """你是 CRM 文档结构化助手。只输出 JSON。把文档中的媒体、联系人、项目与合作记录整理为 records 数组，每条记录只能使用给定标准字段 key。禁止猜测；缺失字段省略。同一媒体在不同合作中可以有多条。每条必须包含 media_name，并给出 source_excerpt（不超过160字）。输出 {\"records\":[...],\"warnings\":[]}，最多100条。"""
    result, usage = deepseek_json_object(system, f"文件：{filename}\n标准字段：{json.dumps(field_help, ensure_ascii=False)}\n\n文档内容：\n{text}", max_tokens=7000)
    records = result.get("records") if isinstance(result.get("records"), list) else []
    clean = []
    for index, record in enumerate(records[:100], 1):
        if not isinstance(record, dict):
            continue
        item = {key: record.get(key) for key in FIELD_KEYS if key in record}
        item["_row_number"] = index
        item["_source_excerpt"] = _text(record.get("source_excerpt"))
        clean.append(item)
    return clean, {"usage": usage, "warnings": [str(item)[:240] for item in (result.get("warnings") or [])[:20]]}


def _normalize_rows(rows: list[dict[str, Any]], mapping: dict[str, str] | None = None) -> list[dict[str, Any]]:
    normalized = []
    for index, row in enumerate(rows, 2):
        if mapping is None:
            item = {HEADER_TO_KEY[header]: row.get(header) for header in HEADER_TO_KEY if header in row}
        elif all(key in FIELD_KEYS or key.startswith("_") for key in row):
            item = {key: row.get(key) for key in FIELD_KEYS if key in row}
        else:
            item = {key: row.get(header) for header, key in mapping.items()}
        item["_row_number"] = row.get("_row_number", index)
        if row.get("_source_excerpt"):
            item["_source_excerpt"] = row["_source_excerpt"]
        for key in list(item):
            if key in DATE_KEYS:
                item[key] = _date(item[key])
            elif key in NUMBER_KEYS:
                item[key] = _number(item[key])
            elif not key.startswith("_"):
                item[key] = _text(item[key])
        if any(value not in (None, "") for key, value in item.items() if not key.startswith("_")):
            normalized.append(item)
    return normalized


def _media_matches(db: Session, row: dict[str, Any]) -> list[Media]:
    if row.get("profile_url"):
        return find_media_identity_matches(db, row["profile_url"])
    name = row.get("media_name")
    if not name:
        return []
    query = db.query(Media).filter(Media.name == name)
    if row.get("country"):
        normalized = normalize_media_payload({"name": name, "country": row["country"]})
        query = query.filter(Media.country == normalized.get("country"))
    if row.get("platform"):
        normalized = normalize_media_payload({"name": name, "platform_type": row["platform"]})
        query = query.filter(Media.platform_type == normalized.get("platform_type"))
    return query.all()


def _preview_rows(db: Session, rows: list[dict[str, Any]]) -> dict[str, Any]:
    previews = []
    for row in rows:
        warnings = []
        if not row.get("media_name"):
            warnings.append("缺少媒体名称，无法写入")
        matches = _media_matches(db, row)
        if len(matches) > 1:
            warnings.append("媒体身份匹配到多条记录，请先合并")
        if row.get("execution_status") and row["execution_status"] not in EXECUTION_STATUSES:
            warnings.append(f"执行状态“{row['execution_status']}”不在标准字典中")
        if row.get("project_code") and row.get("project_name"):
            project = db.query(Project).filter(Project.project_code == row["project_code"]).first()
            if project and project.name != row["project_name"]:
                warnings.append("项目编号已存在，但项目名称不同")
        if not row.get("profile_url") and not row.get("contact_email"):
            warnings.append("缺少主页和联系人邮箱，将按媒体名称匹配")
        action = "冲突" if any("无法写入" in item or "匹配到多条" in item or "名称不同" in item or "不在标准" in item for item in warnings) else ("更新" if len(matches) == 1 else "新增")
        previews.append({
            "row_number": row.get("_row_number"), "import_action": action, "media_name": row.get("media_name"),
            "project_code": row.get("project_code"), "project_name": row.get("project_name"),
            "execution_status": row.get("execution_status"), "product_bundle": row.get("product_models"),
            "tracking_number": row.get("tracking_number"), "warnings": warnings, "source_excerpt": row.get("_source_excerpt"),
        })
    return {
        "total": len(previews), "rows": previews,
        "created_count": sum(item["import_action"] == "新增" for item in previews),
        "updated_count": sum(item["import_action"] == "更新" for item in previews),
        "unchanged_count": 0,
        "conflict_count": sum(item["import_action"] == "冲突" for item in previews),
        "warning_count": sum(len(item["warnings"]) for item in previews),
    }


def create_universal_preview(db: Session, content: bytes, filename: str, user_id: int) -> dict[str, Any]:
    extension = filename.rsplit(".", 1)[-1].casefold() if "." in filename else ""
    if extension not in TABLE_EXTENSIONS | DOCUMENT_EXTENSIONS:
        raise ValueError("支持 CSV、XLSX、PDF、DOCX、TXT 和 Markdown 文件")
    mapping: dict[str, str] = {}
    metadata: dict[str, Any] = {"warnings": []}
    parser = "local_standard"
    if extension in TABLE_EXTENSIONS:
        headers, source_rows = _load_table(content, extension)
        if "媒体名称*" in headers:
            rows = _normalize_rows(source_rows)
        else:
            mapping, metadata = _agent_table_mapping(headers, source_rows)
            if "media_name" not in mapping.values():
                raise ValueError("Agent 未找到媒体名称列，请改用标准模板或调整表头")
            rows = _normalize_rows(source_rows, mapping)
            parser = "agent_mapping"
    else:
        extracted = _extract_document(content, extension)
        source_rows, metadata = _agent_document_records(extracted, filename)
        rows = _normalize_rows(source_rows, {})
        parser = "agent_document"
    preview = _preview_rows(db, rows)
    proposal = {"parser": parser, "mapping": mapping, "records": rows, "preview": preview, "warnings": metadata.get("warnings", [])}
    run = AgentRun(task_type="universal_import", input_type=extension, source_label=filename, source_hash=hashlib.sha256(content).hexdigest(), status="proposed", model=DEEPSEEK_MODEL if parser.startswith("agent") else "本地标准模板", proposal_json=json.dumps(proposal, ensure_ascii=False), usage_json=json.dumps(metadata.get("usage") or {}, ensure_ascii=False), user_id=user_id)
    db.add(run)
    db.commit()
    return {**preview, "draft_id": run.id, "parser": parser, "mapping": mapping, "agent_warnings": metadata.get("warnings", [])}


def _as_date(value: str | None) -> date | None:
    if not value:
        return None
    try:
        return date.fromisoformat(value)
    except ValueError:
        return None


def _split_products(value: str | None) -> list[str]:
    return [item.strip() for item in re.split(r"[；;|\n]+", value or "") if item.strip()]


def _product_with_undo(db: Session, raw_name: str, actions: list[dict[str, Any]]) -> Product:
    matches = find_product_matches(db, raw_name)
    product = find_or_create_product(db, raw_name, "来自统一导入中心")
    if not matches:
        actions.append({"kind": "create", "entity": "product", "id": product.id})
    return product


def confirm_universal_import(db: Session, draft_id: int, user: User) -> dict[str, Any]:
    run = db.get(AgentRun, draft_id)
    if not run or run.task_type != "universal_import":
        raise ValueError("导入草稿不存在")
    if run.status != "proposed":
        raise ValueError("该导入草稿已处理或失效")
    proposal = json.loads(run.proposal_json or "{}")
    rows = proposal.get("records") or []
    preview = _preview_rows(db, rows)
    if preview["conflict_count"]:
        raise ValueError("仍有冲突记录，请修正源文件后重新生成预览")
    actions: list[dict[str, Any]] = []
    batch = ImportBatch(import_type="universal", filename=run.source_label, source_hash=run.source_hash, user_id=user.id, status="processing")
    db.add(batch)
    db.flush()
    created = updated = unchanged = 0
    for row in rows:
        matches = _media_matches(db, row)
        media = matches[0] if matches else None
        parser = proposal.get("parser", "")
        verification_status = "待核验" if parser == "agent_document" else ("部分核验" if parser == "agent_mapping" else "已核验")
        media_values = normalize_media_payload({"name": row.get("media_name"), "country": row.get("country"), "platform_type": row.get("platform"), "category": row.get("category"), "website_url": row.get("profile_url"), "followers_or_traffic": row.get("followers_k"), "cooperation_status": row.get("cooperation_status"), "notes": row.get("notes"), "verification_status": verification_status})
        media_values.update({"data_source": f"统一导入：{run.source_label}", "data_capture_method": "agent_import" if proposal.get("parser", "").startswith("agent") else "import"})
        row_changed = False
        if not media:
            media = Media(**media_values)
            db.add(media); db.flush(); actions.append({"kind": "create", "entity": "media", "id": media.id}); row_changed = True
        else:
            row_changed = set_with_undo(media, "media", media_values, actions)
        if any(row.get(key) for key in ("contact_name", "contact_email", "contact_phone", "contact_role")):
            contact_query = db.query(Contact).filter(Contact.media_id == media.id)
            if row.get("contact_email"):
                contact = contact_query.filter(Contact.email == row.get("contact_email")).first()
            elif row.get("contact_phone"):
                contact = contact_query.filter(Contact.phone == row.get("contact_phone")).first()
            else:
                contact = contact_query.filter(Contact.name == row.get("contact_name"), Contact.role == row.get("contact_role")).first()
            contact_values = {"name": row.get("contact_name"), "role": row.get("contact_role"), "email": row.get("contact_email"), "phone": row.get("contact_phone"), "data_source": f"统一导入：{run.source_label}", "data_capture_method": "import", "is_primary": True}
            if not contact:
                contact = Contact(media_id=media.id, **contact_values); db.add(contact); db.flush(); actions.append({"kind": "create", "entity": "contact", "id": contact.id}); row_changed = True
            else:
                row_changed = set_with_undo(contact, "contact", contact_values, actions) or row_changed
        project = None
        if row.get("project_code") or row.get("project_name"):
            project = db.query(Project).filter(Project.project_code == row.get("project_code")).first() if row.get("project_code") else db.query(Project).filter(Project.name == row.get("project_name")).first()
            if not project:
                project = Project(name=row.get("project_name") or f"导入项目 {row['project_code']}", project_code=row.get("project_code"), status="Active")
                db.add(project); db.flush(); actions.append({"kind": "create", "entity": "project", "id": project.id}); row_changed = True
        campaign = None
        campaign_fields_present = any(row.get(key) for key in ("project_code", "project_name", "execution_status", "collaboration_type", "product_models", "tracking_number", "content_url", "planned_cost", "actual_cost"))
        if campaign_fields_present:
            query = db.query(Campaign).filter(Campaign.media_id == media.id)
            campaign = query.filter(Campaign.project_id == project.id).first() if project else query.filter(Campaign.project_id.is_(None)).first()
            owner = db.query(User).filter(User.email == row.get("owner_email")).first() if row.get("owner_email") else None
            product_names = _split_products(row.get("product_models"))
            product = _product_with_undo(db, product_names[0], actions) if product_names else None
            values = {"project_id": project.id if project else None, "product_id": product.id if product else None, "owner_id": owner.id if owner else None, "collaboration_type": row.get("collaboration_type"), "execution_status": row.get("execution_status") or "待确认", "expected_publish_date": _as_date(row.get("expected_publish_date")), "next_action": row.get("next_action"), "follow_up_date": _as_date(row.get("follow_up_date")), "notes": row.get("notes")}
            if not campaign:
                campaign = Campaign(media_id=media.id, **values); db.add(campaign); db.flush(); actions.append({"kind": "create", "entity": "campaign", "id": campaign.id}); db.add(CampaignStageEvent(campaign_id=campaign.id, user_id=user.id, from_status=None, to_status=campaign.execution_status, action="import", reason="统一导入中心")); row_changed = True
            else:
                old_status = campaign.execution_status
                changed = set_with_undo(campaign, "campaign", values, actions)
                if changed and campaign.execution_status != old_status:
                    db.add(CampaignStageEvent(campaign_id=campaign.id, user_id=user.id, from_status=old_status, to_status=campaign.execution_status, action="import", reason="统一导入中心"))
                row_changed = changed or row_changed
            for product_name in product_names:
                item = _product_with_undo(db, product_name, actions)
                if project:
                    before = db.query(ProjectProduct).filter(ProjectProduct.project_id == project.id, ProjectProduct.product_id == item.id).first()
                    ensure_project_link(db, project.id, item.id)
                    if not before:
                        db.flush(); link = db.query(ProjectProduct).filter(ProjectProduct.project_id == project.id, ProjectProduct.product_id == item.id).first(); actions.append({"kind": "create", "entity": "project_product", "id": link.id}); row_changed = True
            if row.get("tracking_number") or row.get("oa_pi_number"):
                shipment_query = db.query(Shipment).filter(Shipment.campaign_id == campaign.id)
                shipment = shipment_query.filter(Shipment.tracking_number == row.get("tracking_number")).first() if row.get("tracking_number") else shipment_query.filter(Shipment.oa_pi_number == row.get("oa_pi_number")).first()
                if not shipment:
                    shipment = Shipment(campaign_id=campaign.id, oa_pi_number=row.get("oa_pi_number"), carrier=row.get("carrier"), tracking_number=row.get("tracking_number"), shipped_at=_as_date(row.get("shipped_at")), delivered_at=_as_date(row.get("delivered_at")), status="已签收" if row.get("delivered_at") else ("运输中" if row.get("tracking_number") else "待发货")); db.add(shipment); db.flush(); actions.append({"kind": "create", "entity": "shipment", "id": shipment.id}); row_changed = True
                    for product_name in product_names:
                        product_item = _product_with_undo(db, product_name, actions); shipment_item = ShipmentItem(shipment_id=shipment.id, product_id=product_item.id, product_name=product_item.model, quantity=1); db.add(shipment_item); db.flush(); actions.append({"kind": "create", "entity": "shipment_item", "id": shipment_item.id})
            if row.get("content_url"):
                deliverable = db.query(Deliverable).filter(Deliverable.url == row["content_url"]).first()
                if not deliverable:
                    deliverable = Deliverable(campaign_id=campaign.id, deliverable_type=row.get("collaboration_type") or "Other", url=row["content_url"], published_at=_as_date(row.get("published_at"))); db.add(deliverable); db.flush(); actions.append({"kind": "create", "entity": "deliverable", "id": deliverable.id}); row_changed = True
            if row.get("planned_cost") is not None or row.get("actual_cost") is not None:
                reference = f"统一导入：{run.source_label}：第 {row.get('_row_number')} 行"
                cost = db.query(CostItem).filter(CostItem.campaign_id == campaign.id, CostItem.reference_note == reference).first()
                if not cost:
                    cost = CostItem(campaign_id=campaign.id, cost_type="合作费用", planned_amount=row.get("planned_cost"), actual_amount=row.get("actual_cost"), currency=row.get("currency") or "CNY", payment_status=row.get("payment_status") or "未付款", reference_note=reference); db.add(cost); db.flush(); actions.append({"kind": "create", "entity": "cost_item", "id": cost.id}); row_changed = True
        if row_changed:
            if matches:
                updated += 1
            else:
                created += 1
        else:
            unchanged += 1
    summary = {"success_count": len(rows), "created_count": created, "updated_count": updated, "unchanged_count": unchanged, "conflict_count": 0}
    batch.status = "completed"; batch.summary_json = json.dumps(summary, ensure_ascii=False); batch.undo_json = json.dumps(actions, ensure_ascii=False, default=str)
    run.status = "applied"; run.reviewed_by_id = user.id; run.reviewed_at = datetime.utcnow()
    db.commit()
    return {**summary, "batch_id": batch.id, "draft_id": run.id}
